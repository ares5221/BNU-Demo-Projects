#!/usr/bin/env python
# _*_ coding:utf-8 _*_

# 清理data/clean_docx/source中docx文件中的不规则字符，去除图片，并存储在target
# pro版本仅读取内容，不考虑标题信息
import os
from docx import Document
from docx.oxml.ns import qn
import time
import re
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
import docx
from docx.document import Document


def clean_docx():
    root_path = './../../'
    print(os.path.abspath(root_path))
    docx_path = os.path.join(root_path, 'raw_data/book_data/')
    print('待处理文本列表：',os.listdir(docx_path))
    for docx_name in os.listdir(docx_path):
        if docx_name == 'book-52.docx': # tetsttsetestestsetestesteste
            curr_path = os.path.join(docx_path, docx_name)
            print('读取docx文件，注意其中内容可能是表格格式，区分读取。。list存储')
            docx_content = read_word(curr_path)
            print('清理读取的内容中的乱码...')
            clean_content_text = clean_text(docx_content)
            print('保存文件内容...')
            save_name = 'clean_docx222.docx'
            save_target_file(save_name, clean_content_text)


def iter_block_items(parent):
    """
    判断当前block是段落文本还是表格类型
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def read_table(table):
    return [[cell.text for cell in row.cells] for row in table.rows]


def read_word(word_path):
    doc = docx.Document(word_path)
    docx_content = []
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            # print("text", [block.text])
            docx_content.append([block.text])
        elif isinstance(block, Table):
            # print("table", read_table(block))
            tables = read_table(block)
            for tb in tables:
                docx_content.append(tb)
    # print(docx_content)

    return docx_content


def save_target_file(file_name, content):
    save_path = './'
    file_path = os.path.join(save_path, file_name)
    # if not os.path.exists(file_path):
    if True:
        if file_name[-5:] == '.docx':
            save_docx(file_path, content)
        # todo 处理其他类型文件
        if file_name[-5:] == 'pptx':
            pass
    else:
        print(file_name, '该文件已经存在')


def save_docx(path, content):
    document = docx.Document()
    split_content = content.split('\n')
    print('sssssssssss',split_content)
    for con in split_content:
        if con:
            paragraph = document.add_paragraph(con)
            # 设置字体
            document.styles['Normal'].font.name = u'宋体'
            document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            # 设置字体 11字体大小
            document.styles['Normal'].font.size = 140000
    document.save(path)
    curr_time = time.strftime("%Y-%m-%d %H:%M:%S")
    print(path, ' 文件保存完成', curr_time)


def clean_text(contents):
    '''
    清理读取docx内容中的乱码
    :param content:
    :return:
    '''
    curr_content_str = ''
    for content in contents:
        if content:
            is_end = content[0].endswith('。') or content[0].endswith('．') \
                     or content[0].endswith('？') or content[0].endswith('?')\
                     or content[0].endswith('！') or content[0].endswith('!') \
                     or content[0].endswith('〗') or content[0].endswith('"') \
                     or content[0].endswith('：') or content[0].endswith(':') \
                     or content[0].endswith('.')

            if is_end:
                p_text = content[0] + '\n'
            else:
                p_text = content[0]
            curr_content_str += p_text

    # print('!!!!!!!!!!!!',curr_content_str)
    if curr_content_str:
        clean_text = clean_string(curr_content_str)
    return clean_text


def clean_string(to_clean_str):
    IRREGULAR_CHAR = ['|', 'I', '丨', '《', '》', ' ', '[', ']', '【', '】', '\t', '．', '、，丿', '一，〕', '／。0']
    pianpangbushou = ['丨', '亅', '丿', '乛', '乙', '丶', '乚', '十', '厂', '匚', '刂', '卜', '冂', '亻', '八', '勹', '匕', '几', '亠', '冫丷', '冖', '讠', '凵', '卩', '阝', '厶', '廴', '艹', '屮', '彳', '巛', '辶', '彑', '廾', '彐', '宀', '女', '犭', '彡', '尸', '饣', '扌', '氵', '纟', '巳', '兀', '忄', '幺', '弋', '尢', '夂']
    # todo 待补充
    wrong_chinese_word = ['競', '陸','瘛','鑾']

    # todo 省略号转换出错的情况
    wrong_shengluehao = ['“\n', '“““', '一．．', '“，一','“，一', '一“”', '丿、、','']
    # 去除可能出现的乱码字符
    for i_c in IRREGULAR_CHAR:
        if i_c in to_clean_str:
            to_clean_str = to_clean_str.replace(i_c, '')
    for wslh in wrong_shengluehao:
        if wslh in to_clean_str:
            to_clean_str = to_clean_str.replace(wslh, '')
    # 模式 汉字换行汉字 替换为空
    if '|\n' in to_clean_str:
        to_clean_str = to_clean_str.replace('|\n', '')
    if '丨\n' in to_clean_str:
        to_clean_str = to_clean_str.replace('丨\n', '')
    if '》\n' in to_clean_str:
        to_clean_str = to_clean_str.replace('》\n', '')
    if '《\n' in to_clean_str:
        to_clean_str = to_clean_str.replace('《\n', '')
    if '廾\n' in to_clean_str:
        to_clean_str = to_clean_str.replace('廾\n', '')
    if '\n\n' in to_clean_str:
        to_clean_str = to_clean_str.replace('\n\n', '\n')

    pattern_rule = re.compile("\n[\d]+\n")
    if pattern_rule.findall(to_clean_str):
        to_clean_str = re.sub(pattern_rule, '\n', to_clean_str)

    pattern_rule1 = re.compile("\n.{1,3}\n") # '\nI\n','\n]\n',
    if pattern_rule1.findall(to_clean_str):
        to_clean_str = re.sub(pattern_rule1, '', to_clean_str)

    # # 去除换行造成的段落中断情况
    # print('@@@@@', to_clean_str)
    # pattern_rule2 = re.compile("[\u4e00-\u9fa5]\r[\u4e00-\u9fa5]")
    # if pattern_rule2.findall(to_clean_str):
    #     print('curr #############', pattern_rule2.findall(to_clean_str))
    #     # to_clean_str = re.sub(pattern_rule1, '', to_clean_str)
    # 去除可能出现的偏旁部首
    for pianpang in pianpangbushou:
        if pianpang in to_clean_str:
            to_clean_str = to_clean_str.replace(pianpang, '')

    # 去除可能出现的错字
    for wcw in wrong_chinese_word:
        if wcw in to_clean_str:
            to_clean_str = to_clean_str.replace(wcw, '')

    clean_str1 = to_clean_str
    return clean_str1



if __name__ == '__main__':
    is_test = True
    if is_test:
        clean_docx()
    else:
        pass